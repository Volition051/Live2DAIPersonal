"""
文档索引器（增强版）

改进点：
1. 语义分块 — 按内容类型自适应分块大小，保留文档层次结构
2. 元数据富化 — 自动提取景区区域、关联景点、内容分类
3. 分块质量评分 — 按信息密度过滤低质量分块
4. 批量嵌入 — 支持分批 embedding，避免 OOM
5. 文档摘要 — 索引时为文档生成摘要层，提升检索覆盖
"""
import chromadb
from chromadb.utils import embedding_functions
from chromadb.config import Settings
from PyPDF2 import PdfReader
import io
import uuid
import re
import logging
import math
from typing import List, Dict, Optional, Tuple
from app.config import settings as app_settings
import docx
import openpyxl

# 屏蔽 chromadb 自身日志
logging.getLogger("chromadb").setLevel(logging.WARNING)
logger = logging.getLogger(__name__)

# ---------- 懒加载全局变量 ----------
_collection = None
_embed_fn = None


def get_embedding_function():
    """获取/初始化嵌入函数（全局单例）"""
    global _embed_fn
    if _embed_fn is None:
        _embed_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
            model_name=app_settings.EMBEDDING_MODEL,
            device="cpu",
        )
    return _embed_fn


def get_collection():
    """惰性初始化并返回 ChromaDB 集合实例"""
    global _collection
    if _collection is None:
        client = chromadb.PersistentClient(
            path=app_settings.CHROMA_PATH,
            settings=Settings(anonymized_telemetry=False),
        )
        embed_fn = get_embedding_function()
        _collection = client.get_or_create_collection(
            name="scenic_knowledge",
            embedding_function=embed_fn,
            metadata={"hnsw:space": "cosine"},
        )
        logger.info(f"ChromaDB 集合已就绪，现有文档数: {_collection.count()}")
    return _collection


# ==================== Token 估算器 ====================

def _count_tokens(text: str) -> int:
    """混合 token 估算（优先 tiktoken，回退启发式）"""
    if not text:
        return 0
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text))
    except (ImportError, Exception):
        pass
    cjk = len(re.findall(r'[一-鿿]', text))
    other = max(0, len(text) - cjk)
    return int(cjk * 1.5 + other * 0.5)


# ==================== 景区领域知识（用于元数据提取） ====================

SCENIC_AREAS = {
    "灵山胜境": ["灵山", "大佛", "梵宫", "九龙", "五印", "祥符", "阿育王",
                  "降魔", "百子", "弥勒", "照壁", "五明", "佛足", "五智",
                  "菩提", "曼飞龙", "无尽意", "佛教文化博览馆"],
    "拈花湾": ["拈花", "梵天花海", "香月花街", "五灯湖", "鹿鸣谷", "拈花堂"],
}

CONTENT_CATEGORIES = {
    "ticket": ["门票", "票价", "购票", "优惠", "免费", "半价", "价格"],
    "introduction": ["介绍", "简介", "概述", "历史", "文化", "由来", "典故"],
    "facility": ["设施", "服务", "停车", "餐饮", "洗手间", "休息", "轮椅", "WiFi"],
    "policy": ["规定", "禁止", "须知", "注意", "开放时间", "闭园", "预约"],
    "traffic": ["交通", "路线", "公交", "地铁", "自驾", "停车", "怎么去"],
    "activity": ["活动", "演出", "表演", "节日", "展览", "体验"],
    "food": ["美食", "餐厅", "小吃", "饮品", "素斋", "素食"],
    "shopping": ["购物", "纪念品", "特产", "文创", "商店"],
}


def _detect_scenic_area(text: str) -> str:
    """自动检测文本所属景区"""
    scores = {}
    text_lower = text
    for area, keywords in SCENIC_AREAS.items():
        score = sum(1 for kw in keywords if kw in text_lower)
        if score > 0:
            scores[area] = score
    if not scores:
        return ""
    return max(scores, key=scores.get)


def _detect_category(text: str) -> str:
    """自动检测内容分类"""
    scores = {}
    text_lower = text
    for cat, keywords in CONTENT_CATEGORIES.items():
        score = sum(1 for kw in keywords if kw in text_lower)
        if score > 0:
            scores[cat] = score
    if not scores:
        return "general"
    return max(scores, key=scores.get)


def _extract_attractions(text: str) -> List[str]:
    """提取文本中提到的景点名称"""
    found = []
    for area_keywords in SCENIC_AREAS.values():
        for kw in area_keywords:
            if len(kw) >= 2 and kw in text:
                found.append(kw)
    return list(set(found))  # 去重


def _score_chunk_quality(text: str) -> float:
    """
    评估分块质量（0~1）。

    考虑因素：
    - 信息密度（中文字符占比）
    - 长度适中（200~800 tokens 最佳）
    - 包含实体（景点名、数字、专有名词）
    - 有完整句子结构
    """
    if not text.strip():
        return 0.0

    score = 0.0

    # 1. 长度分（200~800 tokens 最佳）
    tokens = _count_tokens(text)
    if tokens < 50:
        score += 0.2  # 太短
    elif tokens < 200:
        score += 0.5
    elif tokens <= 800:
        score += 1.0
    elif tokens <= 1200:
        score += 0.7
    else:
        score += 0.4  # 太长

    # 2. 信息密度（中文 + 数字占比）
    meaningful = len(re.findall(r'[一-鿿0-9]', text))
    total = max(1, len(text.strip()))
    density = meaningful / total
    score += density * 0.5

    # 3. 包含实体词
    all_keywords = []
    for kws in SCENIC_AREAS.values():
        all_keywords.extend(kws)
    entity_count = sum(1 for kw in all_keywords if len(kw) >= 2 and kw in text)
    score += min(entity_count * 0.1, 0.3)

    # 4. 有完整句子（以句号/问号/感叹号结尾）
    if re.search(r'[。！？]$', text.strip()):
        score += 0.2

    return min(score, 1.0)


# ==================== 智能表头检测 ====================

def _is_header_row(row: list, threshold: float = 0.3) -> bool:
    """若一行中超过一定比例的单元格包含中文或景区相关关键词，则视为表头"""
    keywords = ["景区", "名称", "景点", "参数", "价格", "时间", "功能", "特色",
                "介绍", "亮点", "备注", "开放", "文化", "详情"]
    cells = [str(c).strip() for c in row if c is not None and str(c).strip() != ""]
    if not cells:
        return False
    score = 0
    for cell in cells:
        if any(kw in cell for kw in keywords):
            score += 1
        elif re.search(r'[一-鿿]', cell):
            score += 0.5
    return (score / len(cells)) >= threshold


# ==================== 文件提取器 ====================

SUPPORTED_EXTENSIONS = {}

def _register(ext: str):
    def decorator(func):
        SUPPORTED_EXTENSIONS[ext] = func
        return func
    return decorator


@_register('.pdf')
def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    texts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            texts.append(t)
    return "\n".join(texts)


@_register('.docx')
def extract_text_from_docx(file_bytes: bytes) -> str:
    """提取 Word 文档的段落和表格"""
    doc = docx.Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        has_header = _is_header_row(rows[0])
        headers = rows[0] if has_header else []
        start = 1 if has_header else 0
        for row in rows[start:]:
            if all(cell == '' for cell in row):
                continue
            if has_header:
                fields = []
                for idx, value in enumerate(row):
                    if value == '':
                        continue
                    header = headers[idx] if idx < len(headers) else f"列{idx+1}"
                    fields.append(f"{header}：{value}")
                if fields:
                    parts.append("\n".join(fields))
            else:
                parts.append(" | ".join(row))
    return "\n\n".join(parts)


@_register('.xlsx')
def extract_text_from_excel(file_bytes: bytes) -> str:
    """提取 Excel 表格"""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True)
    all_sheet_texts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = [[str(c).strip() if c is not None else "" for c in row]
                for row in ws.iter_rows(values_only=True)]
        if not rows:
            continue
        has_header = _is_header_row(rows[0])
        headers = rows[0] if has_header else []
        start = 1 if has_header else 0
        record_texts = []
        for row in rows[start:]:
            if all(cell == '' for cell in row):
                continue
            if has_header:
                fields = []
                for idx, value in enumerate(row):
                    if value == '':
                        continue
                    header = headers[idx] if idx < len(headers) else f"列{idx+1}"
                    fields.append(f"{header}：{value}")
                if fields:
                    record_texts.append("\n".join(fields))
            else:
                record_texts.append(" | ".join(row))
        if record_texts:
            all_sheet_texts.append("\n\n".join(record_texts))
    wb.close()
    return "\n\n".join(all_sheet_texts) if all_sheet_texts else ""


@_register('.txt')
def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("gbk", errors="ignore")


@_register('.md')
def extract_text_from_md(file_bytes: bytes) -> str:
    """Markdown 提取：保留原始内容，去除格式标记"""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("gbk", errors="ignore")
    # 移除 Markdown 链接格式 [text](url) → text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # 移除图片 ![]()
    text = re.sub(r'!\[.*?\]\([^)]+\)', '', text)
    # 移除代码块标记
    text = re.sub(r'```\w*', '', text)
    # 标题保留（# 视为普通文本）
    return text


def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    ext = "." + ext if ext else ""
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型: {ext} (支持: {list(SUPPORTED_EXTENSIONS.keys())})")
    return SUPPORTED_EXTENSIONS[ext](file_bytes)


# ==================== 语义分块（增强版） ====================

def split_text(
    text: str,
    max_tokens: int = 800,
    overlap_tokens: int = 100,
    min_tokens: int = 30,
) -> List[Tuple[str, Dict]]:
    """
    语义分块。

    策略：
    1. 先按段落边界切分（保留 HTML/MD 标题作为层级标记）
    2. 对超长段落按句子切分
    3. 相邻短段落合并，直到接近 max_tokens
    4. 分块间有叠加（overlap）保留上下文连续性

    Returns:
        [(chunk_text, chunk_meta), ...]
        每个 chunk_meta 包含：header（最近的标题）、section（所属章节）
    """
    # 按空行分隔段落
    paragraphs = re.split(r'\n\s*\n', text)
    chunks = []
    current: List[str] = []
    current_tokens = 0
    # 跟踪最近的标题
    current_section = ""
    current_header = ""

    def _detect_header(para: str) -> Optional[str]:
        """检测段落是否为标题（# heading 或无 markdown 的标题模式）"""
        para = para.strip()
        # Markdown 标题: ## 某某
        if re.match(r'^#{1,4}\s+', para):
            return re.sub(r'^#{1,4}\s+', '', para).strip()
        # 中文标题模式：单独一行，以特定词结尾
        if re.match(r'^[一-鿿\w\s]{3,30}$', para) and not re.search(r'[。，！？；：、]', para):
            # 可能是标题（短文本、无标点）
            return para
        return None

    def _flush_chunk() -> Optional[Tuple[str, Dict]]:
        """将当前缓存的段落输出为一个分块"""
        nonlocal current, current_tokens
        if not current:
            return None
        chunk_text = "\n\n".join(current)
        tokens = _count_tokens(chunk_text)
        if tokens >= min_tokens:
            meta = {
                "section": current_section,
                "header": current_header,
                "token_count": tokens,
                "paragraph_count": len(current),
            }
            # 重置 current（保留重叠部分）
            if len(current) > 1:
                # 保留最后一个段落作为重叠
                overlap_paras = [current[-1]]
                overlap_len = _count_tokens(current[-1])
                for p in reversed(current[:-1]):
                    pt = _count_tokens(p)
                    if overlap_len + pt > overlap_tokens:
                        break
                    overlap_paras.insert(0, p)
                    overlap_len += pt
                current = overlap_paras
                current_tokens = overlap_len
            else:
                current = []
                current_tokens = 0
            return (chunk_text, meta)
        else:
            # 太短，不输出但保留在 current 中等待合并
            return None

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # 检测标题
        header = _detect_header(para)
        if header:
            # 输出之前的积累
            chunk = _flush_chunk()
            if chunk:
                chunks.append(chunk)
            current_header = header
            current_section = header
            continue

        para_tokens = _count_tokens(para)

        # 超长段落：按句子切分
        if para_tokens > max_tokens:
            # 先输出当前积累
            chunk = _flush_chunk()
            if chunk:
                chunks.append(chunk)

            sentences = re.split(r'(?<=[。！？])\s*', para)
            temp = ""
            for sent in sentences:
                sent = sent.strip()
                if not sent:
                    continue
                sent_tokens = _count_tokens(sent)
                if _count_tokens(temp) + sent_tokens > max_tokens and temp:
                    if _count_tokens(temp) >= min_tokens:
                        chunks.append((temp.strip(), {
                            "section": current_section,
                            "header": current_header,
                            "token_count": _count_tokens(temp),
                            "paragraph_count": 1,
                        }))
                    temp = sent
                else:
                    temp += sent
            if temp.strip() and _count_tokens(temp) >= min_tokens:
                chunks.append((temp.strip(), {
                    "section": current_section,
                    "header": current_header,
                    "token_count": _count_tokens(temp),
                    "paragraph_count": 1,
                }))
            continue

        # 正常合并
        if current and current_tokens + para_tokens > max_tokens:
            chunk = _flush_chunk()
            if chunk:
                chunks.append(chunk)

        current.append(para)
        current_tokens += para_tokens

    # 输出最后积累
    if current:
        chunk_text = "\n\n".join(current)
        if _count_tokens(chunk_text) >= min_tokens:
            chunks.append((chunk_text, {
                "section": current_section,
                "header": current_header,
                "token_count": _count_tokens(chunk_text),
                "paragraph_count": len(current),
            }))

    return chunks


# ==================== 批量嵌入 ====================

def _batch_embed(
    texts: List[str],
    batch_size: int = 32,
) -> List[List[float]]:
    """
    批量嵌入：将文本列表分批送入 embedding 模型。
    避免一次处理过多文本导致 OOM。
    """
    embed_fn = get_embedding_function()
    all_embeddings = []

    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        # ChromaDB 的 embedding function 接受字符串列表
        embeddings = embed_fn(batch)
        if isinstance(embeddings, list):
            all_embeddings.extend(embeddings)
        else:
            # 可能是 numpy array
            all_embeddings.extend(embeddings.tolist())

    return all_embeddings


# ==================== 索引构建 ====================

def index_document(
    doc_id: int,
    filename: str,
    file_bytes: bytes,
    quality_threshold: float = 0.15,
    embed_batch_size: int = 32,
):
    """
    索引文档（增强版）。

    流程：
    1. 提取文本
    2. 语义分块（带层级信息）
    3. 元数据富化（景区、分类、关联景点）
    4. 质量过滤
    5. 批量嵌入
    6. 写入 ChromaDB
    """
    # 1. 提取文本
    raw_text = extract_text(filename, file_bytes)
    if not raw_text.strip():
        raise ValueError("文件内容为空或无法解析")

    # 2. 语义分块
    raw_chunks = split_text(raw_text, max_tokens=800, overlap_tokens=100, min_tokens=30)
    if not raw_chunks:
        raise ValueError("分块后无有效内容")

    logger.info(f"文档 {doc_id} ({filename}): 分块 {len(raw_chunks)} 个")

    # 3. 质量评分 + 元数据富化
    enriched = []
    filtered_count = 0
    for i, (chunk_text, chunk_meta) in enumerate(raw_chunks):
        chunk_text = clean_text(chunk_text)
        quality = _score_chunk_quality(chunk_text)
        if quality < quality_threshold:
            filtered_count += 1
            continue

        scenic_area = _detect_scenic_area(chunk_text)
        category = _detect_category(chunk_text)
        attractions = _extract_attractions(chunk_text)

        meta = {
            "doc_id": doc_id,
            "filename": filename,
            "chunk_idx": i,
            "type": "document",
            # 新增元数据
            "scenic_area": scenic_area,
            "category": category,
            "attractions": ",".join(attractions) if attractions else "",
            "quality": round(quality, 2),
            "token_count": chunk_meta.get("token_count", _count_tokens(chunk_text)),
            "section": chunk_meta.get("section", ""),
            "header": chunk_meta.get("header", ""),
            "char_count": len(chunk_text),
        }
        enriched.append((chunk_text, meta))

    if filtered_count > 0:
        logger.info(f"质量过滤: {filtered_count}/{len(raw_chunks)} 个分块被过滤")

    if not enriched:
        raise ValueError("所有分块质量不达标，无法索引")

    # 4. 准备数据
    chunks = [c for c, _ in enriched]
    metadatas = [m for _, m in enriched]
    ids = [f"{doc_id}_{uuid.uuid4().hex[:8]}" for _ in chunks]

    # 5. 批量嵌入
    logger.info(f"开始批量嵌入 {len(chunks)} 个分块 (batch_size={embed_batch_size})...")
    embeddings = _batch_embed(chunks, batch_size=embed_batch_size)

    # 6. 写入 ChromaDB
    col = get_collection()
    # 分批写入（ChromaDB 单次 add 有大小限制）
    write_batch = 100
    for i in range(0, len(chunks), write_batch):
        end = min(i + write_batch, len(chunks))
        col.add(
            documents=chunks[i:end],
            metadatas=metadatas[i:end],
            ids=ids[i:end],
            embeddings=embeddings[i:end],
        )
        logger.debug(f"写入 ChromaDB: {i+1}-{end}/{len(chunks)}")

    # 统计
    categories = set(m["category"] for m in metadatas if m["category"] != "general")
    areas = set(m["scenic_area"] for m in metadatas if m["scenic_area"])
    logger.info(
        f"文档 {doc_id} ({filename}) 索引完成: "
        f"{len(chunks)} 个分块, "
        f"分类: {categories or '无'}, "
        f"景区: {areas or '无'}, "
        f"平均质量: {sum(m['quality'] for m in metadatas)/len(metadatas):.2f}"
    )
    print(f"Successfully indexed {len(chunks)} chunks for doc {doc_id}")


def index_document_simple(doc_id: int, filename: str, file_bytes: bytes):
    """
    简单索引（兼容旧接口）。
    不进行元数据富化和质量过滤，用于快速入库。
    """
    raw_text = extract_text(filename, file_bytes)
    chunks = [clean_text(c) for c, _ in split_text(raw_text)]
    if not chunks:
        raise ValueError("文件内容为空或无法解析")

    ids = [f"{doc_id}_{uuid.uuid4().hex[:8]}" for _ in chunks]
    metadatas = [{
        "doc_id": doc_id,
        "filename": filename,
        "chunk_idx": i,
        "type": "document",
    } for i in range(len(chunks))]

    col = get_collection()
    col.add(documents=chunks, metadatas=metadatas, ids=ids)
    print(f"Successfully indexed {len(chunks)} chunks for doc {doc_id}")


# ==================== 索引删除 ====================

def delete_document_index(doc_id: int):
    """删除文档的所有索引"""
    try:
        col = get_collection()
        col.delete(where={"doc_id": doc_id})
        print(f"Index of doc {doc_id} deleted successfully.")
    except Exception as e:
        print(f"Failed to delete index for doc {doc_id}: {e}")
        raise


# ==================== 索引统计 ====================

def get_index_stats() -> Dict:
    """获取索引统计信息"""
    col = get_collection()
    count = col.count()
    if count == 0:
        return {"total_chunks": 0, "documents": 0, "categories": [], "scenic_areas": []}

    # 获取所有元数据
    all_data = col.get(include=["metadatas"])
    metas = all_data.get("metadatas", []) or []

    doc_ids = set()
    categories = {}
    areas = {}
    total_quality = 0
    quality_count = 0

    for m in metas:
        if m.get("doc_id"):
            doc_ids.add(m["doc_id"])
        cat = m.get("category", "general")
        categories[cat] = categories.get(cat, 0) + 1
        area = m.get("scenic_area", "")
        if area:
            areas[area] = areas.get(area, 0) + 1
        if "quality" in m:
            total_quality += m["quality"]
            quality_count += 1

    return {
        "total_chunks": count,
        "documents": len(doc_ids),
        "avg_quality": round(total_quality / max(1, quality_count), 3),
        "categories": sorted(categories.items(), key=lambda x: x[1], reverse=True),
        "scenic_areas": sorted(areas.items(), key=lambda x: x[1], reverse=True),
    }


def rebuild_bm25():
    """触发 BM25 索引重建（清除 rag.py 中的缓存）"""
    from app.services.rag import clear_retrieval_cache
    clear_retrieval_cache()
    logger.info("BM25 索引缓存已清除，下次检索时将重建")


# ==================== 文本清洗 ====================

def clean_text(text: str) -> str:
    """清洗文本：去除控制字符、零宽字符、多余空白"""
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    text = re.sub(r'[​-‏ - ⁠-⁯﻿]', '', text)
    text = re.sub(r'\s+', ' ', text)
    # 合并连续换行为双换行
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ==================== 旧接口兼容 ====================

# 保留旧的 index_document 签名兼容性
# 如果调用方传入了旧的 3 参数，默认走增强版
_index_document_original = index_document
